from django.shortcuts import render, redirect
from django.urls import reverse
from django.contrib import messages
from django.contrib.sites.shortcuts import get_current_site 
from django.template.loader import render_to_string
from django.core.mail import BadHeaderError, EmailMultiAlternatives
from django.conf import settings
from docx import Document
from django.http import FileResponse, JsonResponse
from django.db.models import Q
from io import BytesIO
from django.utils import timezone
from .decorators import customer_required
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
import random
from django.db.models import F, ExpressionWrapper, IntegerField
from django.db.models.functions import Abs

from .models import *
from user.models import *
from staff.models import *
from .forms import  *
from .mobilenig import *

# Create your views here.

def random_number():
    return int(random.uniform(100000000000000, 999999999999999))
    
def generate_transfer_id():
    while True:
        txn_id = random_number()
        if not Transfer.objects.filter(txn_id=txn_id).exists():
            return txn_id

def generate_bill_id():
    while True:
        txn_id = random_number()
        if not Bill.objects.filter(txn_id=txn_id).exists():
            return txn_id

def generate_topup_id():
    while True:
        txn_id = random_number()
        if not TopUp.objects.filter(txn_id=txn_id).exists():
            return txn_id

def generate_card_id():
    while True:
        card_id = random_number()
        if not Card.objects.filter(card_id=card_id).exists():
            return card_id
        
def random_card_number():
    return int(random.uniform(1000000000000000, 9999999999999999))

def generate_card_number():
    while True:
        card_number = random_card_number()
        if not Card.objects.filter(card_number=card_number).exists():
            return card_number
        
def random_cvv():
    return int(random.uniform(100, 999))

def generate_cvv():
    while True:
        cvv = random_cvv()
        if not Card.objects.filter(cvv=cvv).exists():
            return cvv

def spent(request):
    try:
        txn = Txn.objects.filter(account__id=request.user.id, date__date=datetime.datetime.now())
    except Txn.DoesNotExist:
        txn = None
    trans = [float(tin.amount) for tin in txn]
    return sum(trans)

def account_limit(request):
    return CustomerAccount.objects.get(id=request.user.id).txn_limit

def date_eligible(request):
    return request.user.date_joined + datetime.timedelta(weeks=24)

def transaction_eligible(request):
    if Txn.objects.filter(account__id=request.user.id).exists():
        txn = Txn.objects.filter(account__id=request.user.id)
        trans = [float(tin.amount) for tin in txn]
        return sum(trans)
    else:
        return 0

def expires():
    return datetime.datetime.now().date() + datetime.timedelta(weeks=100)

def get_airtime_status(name):
    service = Provider.objects.get(name=name)
    # Implement logic here to fetch the status for this service
    if service.name == "MTN":
        return MNAirtimeStatus.mtn_get_status()  # Replace with your logic for MTN
    if service.name == "9MOBILE":
        return MNAirtimeStatus.etisalat_get_status()  # Replace with your logic for MTN
    if service.name == "GLO":
        return MNAirtimeStatus.glo_get_status()  # Replace with your logic for MTN
    if service.name == "AIRTEL":
        return MNAirtimeStatus.airtel_get_status()  # Replace with your logic for MTN
    else:
        return "Status Not Available"
    
def get_data_status(name):
    service = Provider.objects.get(name=name)
    # Implement logic here to fetch the status for this service
    if service.name == "MTN":
        return MNDataStatus.mtn_get_status()  # Replace with your logic for MTN
    if service.name == "9MOBILE":
        return MNDataStatus.etisalat_get_status()  # Replace with your logic for MTN
    if service.name == "GLO":
        return MNDataStatus.glo_get_status()  # Replace with your logic for MTN
    if service.name == "AIRTEL":
        return MNDataStatus.airtel_get_status()  # Replace with your logic for MTN
    else:
        return "Status Not Available"

def get_closest_plan(name, target_amount):
    closest_plan = CablePlan.objects.filter(service__name=name).annotate(
        diff=ExpressionWrapper(Abs(F('amount') - target_amount), output_field=IntegerField())
    ).order_by('diff').first()
    return closest_plan

def DataPlansForServiceView(request):
    service_id = request.GET.get("service_id", None)
    if service_id is None or service_id == '':
        # Handle the case when there is no service_id
        return JsonResponse({"error": "No service_id provided"})
    plans = Plan.objects.filter(service_id=service_id)
    plan_data = [{"id": plan.id, "name": plan.name, "amount": plan.amount} for plan in plans]
    return JsonResponse({"plans": plan_data})

def UpdateMTNPlanView():
    plans = MNDataPlan.mtn_get_plan()
    for plan in plans:
        value = int(plan['value'])
        cost = float(plan['cost'])
        f_value = f"{value/1000:.2f} GB"
        f_name = f"MTN {f_value}"
        code = plan['code']
        if Plan.objects.filter(service__name='MTN', name=f_name).exists():
            data = Plan.objects.filter(service__name='MTN', name=f_name)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not Provider.objects.filter(name='MTN').exists():
                new = Provider(name='MTN')
                new.save()
                data = Plan(service=new, name=f_name, amount=cost, product_code=code)
                data.save()
            else:
                obj = Provider.objects.get(name='MTN')
                data = Plan(service=obj, name=f_name, amount=cost, product_code=code)
                data.save()
    
def UpdateEtisalatPlanView():
    plans = MNDataPlan.etisalat_get_plan()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if Plan.objects.filter(service__name='9MOBILE', name=value).exists():
            data = Plan.objects.filter(service__name='9MOBILE', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not Provider.objects.filter(name='9MOBILE').exists():
                new = Provider(name='9MOBILE')
                new.save()
                data = Plan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = Provider.objects.get(name='9MOBILE')
                data = Plan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def UpdateGLOPlanView():
    plans = MNDataPlan.glo_get_plan()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if Plan.objects.filter(service__name='GLO', name=value).exists():
            data = Plan.objects.filter(service__name='GLO', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not Provider.objects.filter(name='GLO').exists():
                new = Provider(name='GLO')
                new.save()
                data = Plan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = Provider.objects.get(name='GLO')
                data = Plan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def UpdateAIRTELPlanView():
    plans = MNDataPlan.airtel_get_plan()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if Plan.objects.filter(service__name='AIRTEL', name=value).exists():
            data = Plan.objects.filter(service__name='AIRTEL', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not Provider.objects.filter(name='AIRTEL').exists():
                new = Provider(name='AIRTEL')
                new.save()
                data = Plan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = Provider.objects.get(name='AIRTEL')
                data = Plan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def CablePlansForServiceView(request):
    service_id = request.GET.get("service_id")
    plans = CablePlan.objects.filter(service_id=service_id)
    plan_data = [{"id": plan.id, "name": plan.name, "amount": plan.amount} for plan in plans]
    return JsonResponse({"plans": plan_data})

def UpdateGOTVPackageView():
    plans = MNCable.gotv_get_package()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if CablePlan.objects.filter(service__name='GOTV', name=value).exists():
            data = CablePlan.objects.filter(service__name='GOTV', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not CableService.objects.filter(name='GOTV').exists():
                new = CableService(name='GOTV')
                new.save()
                data = CablePlan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = CableService.objects.get(name='GOTV')
                data = CablePlan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def UpdateDSTVPackageView():
    plans = MNCable.dstv_get_package()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if CablePlan.objects.filter(service__name='DSTV', name=value).exists():
            data = CablePlan.objects.filter(service__name='DSTV', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not CableService.objects.filter(name='DSTV').exists():
                new = CableService(name='DSTV')
                new.save()
                data = CablePlan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = CableService.objects.get(name='DSTV')
                data = CablePlan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def UpdateStartimesPackageView():
    plans = MNCable.startimes_get_package()
    for plan in plans:
        value = plan['name']
        cost = float(plan['price'])
        code = plan['productCode']
        if CablePlan.objects.filter(service__name='Startimes', name=value).exists():
            data = CablePlan.objects.filter(service__name='Startimes', name=value)
            for datum in data:
                if datum.amount != cost:
                    datum.amount = cost
                    datum.save()
                if datum.product_code != code or not datum.product_code:
                    datum.product_code = code
                    datum.save()
        else:
            if not CableService.objects.filter(name='Startimes').exists():
                new = CableService(name='Startimes')
                new.save()
                data = CablePlan(service=new, name=value, amount=cost, product_code=code)
                data.save()
            else:
                obj = CableService.objects.get(name='Startimes')
                data = CablePlan(service=obj, name=value, amount=cost, product_code=code)
                data.save()
    
def UpdateElectricityBillerView():
    if not ElectricityService.objects.filter(name='Eko Electricity Prepaid').exists():
        new = ElectricityService(name='Eko Electricity Prepaid')
        new.save()
    if not ElectricityService.objects.filter(name='Abuja Electricity Prepaid').exists():
        new = ElectricityService(name='Abuja Electricity Prepaid')
        new.save()
    if not ElectricityService.objects.filter(name='Kaduna Electricity Prepaid').exists():
        new = ElectricityService(name='Kaduna Electricity Prepaid')
        new.save()
    if not ElectricityService.objects.filter(name='Ibadan Electricity Prepaid').exists():
        new = ElectricityService(name='Ibadan Electricity Prepaid')
        new.save()
    if ElectricityService.objects.exclude(name__in=['Eko Electricity Prepaid', 'Abuja Electricity Prepaid', 'Kaduna Electricity Prepaid', 'Ibadan Electricity Prepaid']).exists():
        other_services = ElectricityService.objects.exclude(name__in=['Eko Electricity Prepaid', 'Abuja Electricity Prepaid', 'Kaduna Electricity Prepaid', 'Ibadan Electricity Prepaid'])
        for services in other_services:
            services.delete()
        
@customer_required
def CustomerHomeView(request):
    account = CustomerAccount.objects.get(id=request.user.id)
    obj = Alert.objects.filter(user=request.user, statement=False, modify_limit=False).order_by('-date')[:4]
    obj_c = Alert.objects.filter(user=request.user, read=False, statement=False, modify_limit=False).count()
    obj_co = Alert.objects.filter(user=request.user).count()
    data = Exchange.objects.order_by('-date')[:4]
    data_c = Exchange.objects.order_by('-date').count()
    data_date = Exchange.objects.order_by('-date').first()
    post = News.objects.order_by('-date')[:4]
    post_c = News.objects.exclude(read = request.user).count()
    post_co = News.objects.count()
    message_c = Message.objects.filter(user=request.user, read=False).count()
    if Card.objects.filter(account__id=request.user.id, expires__gt=datetime.datetime.now().date()).exists():
        card = Card.objects.get(account__id=request.user.id, expires__gt=datetime.datetime.now().date())
    else:
        card = False
    context = {
        'account':account,
        'obj':obj,
        'obj_c':obj_c,
        'obj_co':obj_co,
        'data':data,
        'data_c':data_c,
        'data_date':data_date,
        'post':post,
        'post_c':post_c,
        'post_co':post_co,
        'message_c':message_c,
        'card':card,
        }
    return render(request, 'customer/home.html', context)

@customer_required
def MoreAlertView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    obj = Alert.objects.filter(user=request.user, statement=False, modify_limit=False).order_by('-date')
    obj_c = Alert.objects.filter(user=request.user, read=False, statement=False, modify_limit=False).count()
    paginator = Paginator(obj, 10)
    page = request.GET.get('page')
    try:
        obj = paginator.page(page)
    except PageNotAnInteger:
        obj = paginator.page(1)
    except EmptyPage:
        obj = paginator.page(paginator.num_pages)
    context = {
        'obj':obj,
        'obj_c':obj_c,
        'message_c':message_c,
        }
    return render(request, 'customer/morealert.html', context)

@customer_required
def DetailAlertView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    alert = Alert.objects.read_alert(alert_id=pk)
    context = {
        'message_c':message_c,
        'alert':alert
        }
    return render(request, 'customer/detailalert.html', context)

@customer_required
def MoreExchangeView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    data = Exchange.objects.order_by('-date')
    data_date = Exchange.objects.order_by('-date').first()
    paginator = Paginator(data, 10)
    page = request.GET.get('page')
    try:
        data = paginator.page(page)
    except PageNotAnInteger:
        data = paginator.page(1)
    except EmptyPage:
        data = paginator.page(paginator.num_pages)
    context = {
        'message_c':message_c,
        'data':data,
        'data_date':data_date,
        }
    return render(request, 'customer/moreexchange.html', context)

@customer_required
def MoreNewsView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = News.objects.order_by('-date')
    post_c = News.objects.exclude(read = request.user).count()
    paginator = Paginator(post, 10)
    page = request.GET.get('page')
    try:
        post = paginator.page(page)
    except PageNotAnInteger:
        post = paginator.page(1)
    except EmptyPage:
        post = paginator.page(paginator.num_pages)
    context = {
        'message_c':message_c,
        'post':post,
        'post_c':post_c,
        }
    return render(request, 'customer/morenews.html', context)

@customer_required
def DetailNewsView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    news = News.objects.read_news(user_id=request.user, news_id=pk)
    context = {
        'message_c':message_c,
        'news':news
        }
    return render(request, 'customer/detailnews.html', context)

@customer_required
def MoreView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    support_c = Support.objects.filter(customer=request.user, read=False).count()
    return render(request, 'customer/more.html', {'message_c':message_c, 'support_c':support_c})

@customer_required
def TransferView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    obj = Transfer.objects.filter(user=request.user, status='Completed').order_by('-date')[:4]
    obj_c = Transfer.objects.filter(user=request.user, status='Completed').order_by('-date').count()
    txn_limit = account_limit(request) - spent(request)
    return render(request, 'customer/transfer.html', {'message_c':message_c, 'obj':obj, 'obj_c':obj_c, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def MoreTransferView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    transfer = Transfer.objects.filter(user=request.user, status='Completed').order_by('-date')
    paginator = Paginator(transfer, 10)
    page = request.GET.get('page')
    try:
        transfer = paginator.page(page)
    except PageNotAnInteger:
        transfer = paginator.page(1)
    except EmptyPage:
        transfer = paginator.page(paginator.num_pages)
    return render(request, 'customer/moretransfer.html', {'message_c':message_c, 'transfer':transfer})

@customer_required
def DetailTransferView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    transfer = Transfer.objects.get(id=pk)
    context = {
        'message_c':message_c,
        'transfer':transfer
        }
    return render(request, 'customer/detailtransfer.html', context)

@customer_required
def BillView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    saved = Bill.objects.filter(saved=True, account__id=request.user.id).order_by('-date')[:2]
    recent = Bill.objects.filter(status='Completed', account__id=request.user.id).order_by('-date')[:4]
    saved_c = Bill.objects.filter(saved=True, account__id=request.user.id).order_by('-date').count()
    recent_c = Bill.objects.filter(status='Completed', account__id=request.user.id).order_by('-date').count()
    txn_limit = account_limit(request) - spent(request)
    context = {
        'message_c':message_c,
        'saved':saved,
        'recent':recent,
        'saved_c':saved_c,
        'recent_c':recent_c,
        'txn_limit':txn_limit,
        'account_limit':account_limit(request)
        }
    return render(request, 'customer/bill.html', context)

@customer_required
def MoreBillSavedView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    saved = Bill.objects.filter(account__id=request.user.id, saved=True).order_by('-date')
    paginator = Paginator(saved, 10)
    page = request.GET.get('page')
    try:
        saved = paginator.page(page)
    except PageNotAnInteger:
        saved = paginator.page(1)
    except EmptyPage:
        saved = paginator.page(paginator.num_pages)
    return render(request, 'customer/moresavedbill.html', {'message_c':message_c, 'saved':saved})

@customer_required
def MoreBillView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    recent = Bill.objects.filter(account__id=request.user.id, status='Completed').order_by('-date')
    paginator = Paginator(recent, 10)
    page = request.GET.get('page')
    try:
        recent = paginator.page(page)
    except PageNotAnInteger:
        recent = paginator.page(1)
    except EmptyPage:
        recent = paginator.page(paginator.num_pages)
    return render(request, 'customer/morerecentbill.html', {'message_c':message_c, 'recent':recent})

@customer_required
def NewBillView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    txn_limit = account_limit(request) - spent(request)
    return render(request, 'customer/newbill.html', {'message_c':message_c, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def DetailBillView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    bill = Bill.objects.get(id=pk)
    context = {
        'message_c':message_c,
        'bill':bill
        }
    return render(request, 'customer/detailbill.html', context)

@customer_required
def CablePaymentView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    
    form = CableForm(request.POST or None)
    if request.method == 'POST':
        if form.is_valid():
            obj = form.save(commit=False)
            obj.account = post
            obj.save()
            bill = Bill(account=post, txn_id=generate_bill_id(), cable=obj)
            bill.save()
            if bill.cable.service.name == 'GOTV':
                result = MNCable.gotv_get_details(pk=bill.pk)
            if bill.cable.service.name == 'DSTV':
                result = MNCable.dstv_get_details(pk=bill.pk)
            if bill.cable.service.name == 'Startimes':
                result = MNCable.startimes_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Card Number And Try Again.')
                return redirect('cable')
            detail = result
            if bill.cable.plan:
                if bill.cable.service.name == 'Startimes':
                    amount = detail['billAmount']
                else:
                    amount = detail['amount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = bill.cable.plan
                if old_plan == new_plan:
                    messages.error(request, 'You Want To Recharge Thesame Plan. Pick Renew Subscription In Your Form')
                    return redirect('cable')
            if 'saved' in request.POST:
                bill.saved = True
                bill.save()
                messages.success(request, 'Bill Saved')
            messages.info(request, 'Confirm Transaction')
            return redirect(reverse('confirmbill', args=[bill.pk]))
    return render(request, 'customer/cablepayment.html', {'message_c':message_c, 'form':form, 'post':post, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def ElectricityPaymentView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    
    form = ElectricityForm(request.POST or None)
    if request.method == 'POST':
        if form.is_valid():
            obj = form.save(commit=False)
            obj.account = post
            obj.save()
            bill = Bill(account=post, txn_id=generate_bill_id(), electricity=obj)
            bill.save()
            if 'saved' in request.POST:
                bill.saved = True
                bill.save()
                messages.success(request, 'Bill Saved')
            messages.info(request, 'Confirm Transaction')
            return redirect(reverse('confirmbill', args=[bill.pk]))
    return render(request, 'customer/electricitypayment.html', {'message_c':message_c, 'form':form, 'post':post, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def RepeatElectricityView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    repeat = Electricity.objects.get(id=pk)
    bill = Bill.objects.get(electricity=repeat)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        amount = request.POST['amount']
        if repeat.phone_number:
            new = Electricity(account=post, service=repeat.service, amount=amount, meter_number=repeat.meter_number, phone_number=repeat.phone_number)
            new.save()
        else:
            new = Electricity(account=post, service=repeat.service, amount=amount, meter_number=repeat.meter_number)
            new.save()
        bill = Bill(account=post, txn_id=generate_bill_id(), electricity=new)
        bill.save()
        if 'saved' in request.POST:
            bill.saved = True
            bill.save()
            messages.success(request, 'Bill Saved')
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmbill', args=[bill.pk]))
    return render(request, 'customer/repeatelectricity.html', {'message_c':message_c, 'repeatelectricity':repeat, 'post':post, 'bill':bill, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def ConfirmBillView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    bill = Bill.objects.get(id=pk)
    if bill.cable:
        if bill.cable.service.name == 'GOTV':
            result = MNCable.gotv_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Card Number And Try Again.')
                return redirect('cable')
            detail = result
            name = detail['firstName'] + ' ' + detail['lastName']
            customerType = detail['customerType']
            dueDate = detail['dueDate']
            returnCode = None
            accountStatus = detail['accountStatus']
            invoicePeriod = detail['invoicePeriod']
            customerNumber = detail['customerNumber']
            if bill.cable.plan:
                amount = detail['amount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = bill.cable.plan
                old_amount = old_plan.amount
                new_amount = bill.cable.plan.amount
            else:
                amount = detail['amount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = old_plan
                old_amount = old_plan.amount
                new_amount = old_amount
                bill.cable.plan = old_plan
                bill.cable.save()
            bill.cable.customer_name = name
            bill.cable.customer_number = customerNumber
            bill.cable.save()
            bill.amount = new_amount
            bill.save()
        if bill.cable.service.name == 'DSTV':
            result = MNCable.dstv_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Card Number And Try Again.')
                return redirect('cable')
            detail = result
            name = detail['firstName'] + ' ' + detail['lastName']
            customerType = detail['customerType']
            dueDate = detail['dueDate']
            returnCode = None
            accountStatus = detail['accountStatus']
            invoicePeriod = detail['invoicePeriod']
            customerNumber = detail['customerNumber']
            if bill.cable.plan:
                amount = detail['amount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = bill.cable.plan
                old_amount = old_plan.amount
                new_amount = bill.cable.plan.amount
            else:
                amount = detail['amount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = old_plan
                old_amount = old_plan.amount
                new_amount = old_amount
                bill.cable.plan = old_plan
                bill.cable.save()
            bill.cable.customer_name = name
            bill.cable.customer_number = customerNumber
            bill.cable.save()
            bill.amount = new_amount
            bill.save()
        if bill.cable.service.name == 'Startimes':
            result = MNCable.startimes_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Card Number And Try Again.')
                return redirect('cable')
            detail = result
            name = detail['customerName']
            customerType = detail['customerType']
            dueDate = None
            returnCode = detail['returnCode']
            accountStatus = None
            invoicePeriod = None
            customerNumber = detail['customerNumber']
            if bill.cable.plan:
                amount = detail['billAmount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = bill.cable.plan
                old_amount = old_plan.amount
                new_amount = bill.cable.plan.amount
            else:
                amount = detail['billAmount']
                if bill.cable.service.name == 'GOTV':
                    old_plan = get_closest_plan('GOTV', amount)
                if bill.cable.service.name == 'DSTV':
                    old_plan = get_closest_plan('DSTV', amount)
                if bill.cable.service.name == 'Startimes':
                    old_plan = get_closest_plan('Startimes', amount)
                new_plan = old_plan
                old_amount = old_plan.amount
                new_amount = old_amount
                bill.cable.plan = old_plan
                bill.cable.save()
            bill.cable.customer_name = name
            bill.cable.customer_number = customerNumber
            bill.cable.save()
            bill.amount = new_amount
            bill.save()
        if request.method == 'POST':
            if request.user.balance < bill.amount:
                messages.error(request, 'Insufficient Funds')
                return redirect('cable')
            else:
                messages.info(request, 'Last Step')
                return redirect(reverse('validatebill', args=[bill.pk]))
        context = {
            'message_c':message_c,
            'bill':bill,
            'name':name,
            'customerType':customerType,
            'dueDate':dueDate,
            'returnCode':returnCode,
            'accountStatus':accountStatus,
            'invoicePeriod':invoicePeriod,
            'customerNumber':customerNumber,
            'old_amount':old_amount,
            'new_amount':new_amount,
            'old_plan':old_plan,
            'new_plan':new_plan
            }
        return render(request, 'customer/confirmbill.html', context)
    if bill.electricity:
        if bill.electricity.service.name == 'Eko Electricity Prepaid':
            result = MNElectricity.eko_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Meter Number And Try Again.')
                return redirect('electricity')
            detail = result
            name = detail['customerName']
            customerAddress = detail['customerAddress']
            customerDistrict = detail['customerDistrict']
            status = detail['status']
            responseMessage = detail['responseMessage']
            customerReference = None
            minimumVend = None
            tariff = None
            freeUnits = None
            outstandingAmount = None
            amount = bill.electricity.amount
            bill.electricity.customer_name = name
            bill.electricity.customer_address = customerAddress
            bill.electricity.customer_district = customerDistrict
            bill.electricity.save()
            bill.amount = amount
            bill.save()
        if bill.electricity.service.name == 'Abuja Electricity Prepaid':
            result = MNElectricity.abuja_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Meter Number And Try Again.')
                return redirect('electricity')
            detail = result
            name = detail['customerName']
            customerAddress = detail['customerAddress']
            customerDistrict = None
            status = None
            customerReference = detail['customerReference']
            minimumVend = detail['minimumVend']
            tariff = detail['tariff']
            responseMessage = detail['responseMessage']
            freeUnits = None
            outstandingAmount = None
            amount = bill.electricity.amount
            bill.electricity.customer_name = name
            bill.electricity.customer_address = customerAddress
            bill.electricity.customer_reference = customerReference
            bill.electricity.save()
            bill.amount = amount
            bill.save()
        if bill.electricity.service.name == 'Kaduna Electricity Prepaid':
            result = MNElectricity.kaduna_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Meter Number And Try Again.')
                return redirect('electricity')
            detail = result
            name = detail['customerName']
            customerAddress = detail['customerAddress']
            customerDistrict = None
            status = None
            customerReference = None
            minimumVend = detail['minimumAmount']
            tariff = detail['tariff']
            responseMessage = detail['responseMessage']
            freeUnits = detail['freeUnits']
            outstandingAmount = detail['outstandingAmount']
            amount = bill.electricity.amount
            bill.electricity.customer_name = name
            bill.electricity.customer_address = customerAddress
            bill.electricity.tariff_code = tariff
            bill.electricity.save()
            bill.amount = amount
            bill.save()
        if bill.electricity.service.name == 'Ibadan Electricity Prepaid':
            result = MNElectricity.ibadan_get_details(pk=bill.pk)
            if result == 'Error':
                messages.error(request, 'An Error Occurred. Check The Meter Number And Try Again.')
                return redirect('electricity')
            detail = result
            name = detail['customerName']
            customerAddress = detail['customerAddress']
            customerDistrict = None
            status = None
            customerReference = None
            minimumVend = detail['minimumAmount']
            tariff = detail['tariff']
            responseMessage = detail['responseMessage']
            freeUnits = detail['freeUnits']
            outstandingAmount = detail['outstandingAmount']
            amount = bill.electricity.amount
            bill.electricity.customer_name = name
            if customerAddress:
                bill.electricity.customer_address = customerAddress
            bill.electricity.save()
            bill.amount = amount
            bill.save()
        if request.method == 'POST':
            if request.user.balance < bill.amount:
                messages.error(request, 'Insufficient Funds')
                return redirect('electricity')
            else:
                messages.info(request, 'Last Step')
                return redirect(reverse('validatebill', args=[bill.pk]))
        context = {
            'message_c':message_c,
            'bill':bill,
            'name':name,
            'customerAddress':customerAddress,
            'customerDistrict':customerDistrict,
            'status':status,
            'responseMessage':responseMessage,
            'customerReference':customerReference,
            'minimumVend':minimumVend,
            'tariff':tariff,
            'freeUnits':freeUnits,
            'outstandingAmount':outstandingAmount,
            'amount':amount,
            }
        return render(request, 'customer/confirmbill.html', context)

@customer_required
def ValidateBillView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    bill = Bill.objects.get(id=pk)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        if txn_limit < bill.amount:
            messages.error(request, 'Daily Transaction Limit Has Been Reached')
            return redirect('bill')
        else:
            pin = request.POST['pin']
            if bill.account.user.pin == pin:
                if bill.cable:
                    if bill.cable.service.name == 'GOTV':
                        MNCable.gotv_recharge(request, pk=bill.pk)
                    if bill.cable.service.name == 'DSTV':
                        MNCable.dstv_recharge(request, pk=bill.pk)
                    if bill.cable.service.name == 'Startimes':
                        MNCable.startimes_recharge(request, pk=bill.pk)
                if bill.electricity:
                    if bill.electricity.service.name == 'Eko Electricity Prepaid':
                        MNElectricity.eko_recharge(request, pk=bill.pk)
                    if bill.electricity.service.name == 'Abuja Electricity Prepaid':
                        MNElectricity.abuja_recharge(request, pk=bill.pk)
                    if bill.electricity.service.name == 'Kaduna Electricity Prepaid':
                        MNElectricity.kaduna_recharge(request, pk=bill.pk)
                    if bill.electricity.service.name == 'Ibadan Electricity Prepaid':
                        MNElectricity.ibadan_recharge(request, pk=bill.pk)
                return redirect('bill')
            else:
                messages.error(request, 'Invalid Pin')
                return redirect(reverse('validatebill', args=[bill.pk]))
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def TopUpView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    saved = TopUp.objects.filter(saved=True, account__id=request.user.id).order_by('-date')[:2]
    recent = TopUp.objects.filter(status='Completed', account__id=request.user.id).order_by('-date')[:4]
    saved_c = TopUp.objects.filter(saved=True, account__id=request.user.id).order_by('-date').count()
    recent_c = TopUp.objects.filter(status='Completed', account__id=request.user.id).order_by('-date').count()
    txn_limit = account_limit(request) - spent(request)
    context = {
        'message_c':message_c,
        'saved':saved,
        'recent':recent,
        'saved_c':saved_c,
        'recent_c':recent_c, 'txn_limit':txn_limit, 'account_limit':account_limit(request)
        }
    return render(request, 'customer/topup.html', context)

@customer_required
def NewTopUpView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    txn_limit = account_limit(request) - spent(request)
    context = {
        'message_c':message_c,
        'txn_limit':txn_limit,
        'account_limit':account_limit(request)
        }
    return render(request, 'customer/newtopup.html', context)

@customer_required
def DetailTopView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    topup = TopUp.objects.get(id=pk)
    context = {
        'message_c':message_c,
        'topup':topup
        }
    return render(request, 'customer/detailtopup.html', context)

@customer_required
def MoreTopSavedView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    saved = TopUp.objects.filter(account__id=request.user.id, saved=True).order_by('-date')
    paginator = Paginator(saved, 10)
    page = request.GET.get('page')
    try:
        saved = paginator.page(page)
    except PageNotAnInteger:
        saved = paginator.page(1)
    except EmptyPage:
        saved = paginator.page(paginator.num_pages)
    return render(request, 'customer/moresavedtopup.html', {'message_c':message_c, 'savedtop':saved})

@customer_required
def MoreTopView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    recent = TopUp.objects.filter(account__id=request.user.id, status='Completed').order_by('-date')
    paginator = Paginator(recent, 10)
    page = request.GET.get('page')
    try:
        recent = paginator.page(page)
    except PageNotAnInteger:
        recent = paginator.page(1)
    except EmptyPage:
        recent = paginator.page(paginator.num_pages)
    return render(request, 'customer/morerecenttopup.html', {'message_c':message_c, 'recenttop':recent})

@customer_required
def AirtimeView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    
    providers = Provider.objects.all()

    # Create a dictionary to store service statuses
    provider_statuses = {}

    # Loop through the services and fetch their statuses using the custom method
    for provider in providers:
        status = get_airtime_status(provider.name)
        provider_statuses[provider.name] = status

    form = AirtimeForm(request.POST or None, provider_statuses=provider_statuses)
    
    if request.method == 'POST':
        if form.is_valid():
            obj = form.save(commit=False)
            obj.account = post
            if request.user.balance < obj.amount:
                messages.error(request, "Insufficient Funds")
                return redirect('airtime')
            if obj.amount < float(100):
                messages.error(request, 'Minimum Amount Is ₦100')
                return redirect('airtime')
            if len(obj.phone_number) > 11 or len(obj.phone_number) < 11:
                messages.error(request, 'Phone Number Should Be 11 Digits')
                return redirect('airtime')
            obj.save()
            topup = TopUp(account=post, txn_id=generate_topup_id(), airtime=obj, amount=obj.amount, phone_number=obj.phone_number)
            topup.save()
            if 'saved' in request.POST:
                topup.saved = True
                topup.save()
                messages.success(request, 'Top Up Saved')
            messages.info(request, 'Confirm Transaction')
            return redirect(reverse('confirmtop', args=[topup.pk]))
    return render(request, 'customer/airtimepayment.html', {'message_c':message_c, 'form':form, 'post':post, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def RepeatAirtimeView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    repeat = Airtime.objects.get(id=pk)
    topup = TopUp.objects.get(airtime=repeat)
    txn_limit = account_limit(request) - spent(request)
    airtime_status = get_airtime_status(topup.airtime.service.name)
    if request.method == 'POST':
        amount = request.POST['amount']
        phone_number = request.POST['phone_number']
        if request.user.balance < float(amount):
            messages.error(request, "Insufficient Funds")
            return redirect(reverse('repeatairtime', args=[repeat.pk]))
        if float(amount) < float(100):
            messages.error(request, 'Minimum Amount Is ₦100')
            return redirect(reverse('repeatairtime', args=[repeat.pk]))
        if len(phone_number) > 11 or len(phone_number) < 11:
            messages.error(request, 'Phone Number Should Be 11 Digits')
            return redirect(reverse('repeatairtime', args=[repeat.pk]))
        new = Airtime(account=post, service=repeat.service, amount=amount, phone_number=phone_number)
        new.save()
        topup = TopUp(account=post, txn_id=generate_topup_id(), airtime=new, amount=new.amount, phone_number=new.phone_number)
        topup.save()
        if 'saved' in request.POST:
            topup.saved = True
            topup.save()
            messages.success(request, 'Top Up Saved')
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmtop', args=[topup.pk]))
    return render(request, 'customer/repeatairtimepayment.html', {'message_c':message_c, 'repeatairtime':repeat, 'airtime_status':airtime_status, 'post':post, 'topup':topup, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def DataView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    
    providers = Provider.objects.all()
    
    # Create a dictionary to store service statuses
    provider_statuses = {}

    # Loop through the services and fetch their statuses using the custom method
    for provider in providers:
        status = get_data_status(provider.name)
        provider_statuses[provider.id] = (provider.name, status)
    
    if request.method == 'POST':
        form = DataForm(request.POST, provider_statuses=provider_statuses)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.account = post
            if request.user.balance < obj.plan.amount:
                messages.error(request, "Insufficient Funds")
                return redirect('data')
            if len(obj.phone_number) > 11 or len(obj.phone_number) < 11:
                messages.error(request, 'Phone Number Should Be 11 Digits')
                return redirect('data')
            obj.save()
            topup = TopUp(account=post, txn_id=generate_topup_id(), data=obj, amount=obj.plan.amount, phone_number=obj.phone_number)
            topup.save()
            if 'saved' in request.POST:
                topup.saved = True
                topup.save()
                messages.success(request, 'Top Up Saved')
            messages.info(request, 'Confirm Transaction')
            return redirect(reverse('confirmtop', args=[topup.pk]))
    else:
        form = DataForm(provider_statuses=provider_statuses)
    return render(request, 'customer/datapayment.html', {'message_c':message_c, 'form':form, 'post':post, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def RepeatDataView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = CustomerAccount.objects.get(id=request.user.id)
    repeat = Data.objects.get(id=pk)
    topup = TopUp.objects.get(data=repeat)
    txn_limit = account_limit(request) - spent(request)
    data_status = get_data_status(topup.data.service.name)
    if request.method == 'POST':
        phone_number = request.POST['phone_number']
        if request.user.balance < repeat.plan.amount:
            messages.error(request, "Insufficient Funds")
            return redirect(reverse('repeatdata', args=[repeat.pk]))
        if len(phone_number) > 11 or len(phone_number) < 11:
            messages.error(request, 'Phone Number Should Be 11 Digits')
            return redirect(reverse('repeatdata', args=[repeat.pk]))
        new = Data(account=post, service=repeat.service, plan=repeat.plan, phone_number=phone_number)
        new.save()
        topup = TopUp(account=post, txn_id=generate_topup_id(), data=new, amount=new.plan.amount, phone_number=new.phone_number)
        topup.save()
        if 'saved' in request.POST:
            topup.saved = True
            topup.save()
            messages.success(request, 'Top Up Saved')
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmtop', args=[topup.pk]))
    return render(request, 'customer/repeatdatapayment.html', {'message_c':message_c, 'repeatdata':repeat, 'data_status':data_status, 'post':post, 'topup':topup, 'txn_limit':txn_limit, 'account_limit':account_limit(request)})

@customer_required
def ConfirmTopView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    topup = TopUp.objects.get(id=pk)
    if request.method == 'POST':
        messages.info(request, 'Last Step')
        return redirect(reverse('validatetop', args=[topup.pk]))
    context = {
        'message_c':message_c,
        'topup':topup
        }
    return render(request, 'customer/confirmtopup.html', context)

@customer_required
def ValidateTopView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    topup = TopUp.objects.get(id=pk)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        if txn_limit < topup.amount:
            messages.error(request, 'Daily Transaction Limit Has Been Reached')
            return redirect('topup')
        else:
            pin = request.POST['pin']
            if topup.account.user.pin == pin:
                if topup.airtime:
                    if topup.airtime.service.name == 'MTN':
                        MNAirtimeRecharge.mtn_recharge(request, topup_id=pk)
                    elif topup.airtime.service.name == '9MOBILE':
                        MNAirtimeRecharge.etisalat_recharge(request, topup_id=pk)
                    elif topup.airtime.service.name == 'GLO':
                        MNAirtimeRecharge.glo_recharge(request, topup_id=pk)
                    elif topup.airtime.service.name == 'AIRTEL':
                        MNAirtimeRecharge.airtel_recharge(request, topup_id=pk)
                elif topup.data:
                    if topup.data.service.name == 'MTN':
                        MNDataRecharge.mtn_recharge(request, topup_id=pk)
                    elif topup.data.service.name == '9MOBILE':
                        MNDataRecharge.etisalat_recharge(request, topup_id=pk)
                    elif topup.data.service.name == 'GLO':
                        MNDataRecharge.glo_recharge(request, topup_id=pk)
                    elif topup.data.service.name == 'AIRTEL':
                        MNDataRecharge.airtel_recharge(request, topup_id=pk)
                return redirect('topup')
            else:
                messages.error(request, 'Invalid Pin')
                return redirect(reverse('validatetop', args=[topup.pk]))
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def LocalTransferView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    obj = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        receiver = request.POST['receiver']
        amount = request.POST['amount']
        purpose = request.POST['purpose']
        if len(receiver) > 10 or len(receiver) < 10:
            messages.error(request, "Length Of Account Number Should Be 10")
            return redirect('localtransfer')
        if CustomerAccount.objects.filter(account_number=int(receiver)).exists() == False:
            messages.error(request, "Receiver's Account Not Found")
            return redirect('localtransfer')
        if request.user.balance < float(amount):
            messages.error(request, "Insufficient Funds")
            return redirect('localtransfer')
        if float(amount) < float(100):
            messages.error(request, "Minimum Amount Is ₦100")
            return redirect('localtransfer')
        if float(amount) > float(100000):
            messages.error(request, "Maximum Amount Is ₦100 000")
            return redirect('localtransfer')
        receiver_account = CustomerAccount.objects.get(account_number=receiver)
        if receiver_account == obj:
            messages.error(request, 'You Cannot Transfer To Yourself')
            return redirect('localtransfer')
        if 'add' in request.POST:
            if Beneficiary.objects.filter(confirm=True, account=obj, beneficiary=receiver_account).exists():
                messages.error(request, 'Receiver Is Already A Beneficiary')
            else:
                new = Beneficiary(confirm=True, account=obj, beneficiary=receiver_account)
                new.save()
                messages.success(request, 'Receiver Successfully Added To Beneficiaries')
        post = Transfer(sender=obj, txn_id=generate_transfer_id(), receiver=receiver_account, amount=float(amount), purpose=purpose)
        post.save()
        post.user.add(request.user)
        post.user.add(receiver_account.user)
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmtransfer', args=[post.pk]))
    context = {
        'message_c':message_c,
        'obj':obj,
        'txn_limit':txn_limit,
        'account_limit':account_limit(request),
        }
    return render(request, 'customer/localtransfer.html', context)

@customer_required
def RepeatTransferView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    transfer = Transfer.objects.get(id=pk)
    obj = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        receiver = request.POST['receiver']
        amount = request.POST['amount']
        purpose = request.POST['purpose']
        if len(receiver) > 10 or len(receiver) < 10:
            messages.error(request, "Length Of Account Number Should Be 10")
            return redirect(reverse('repeattransfer', args=[transfer.pk]))
        if CustomerAccount.objects.filter(account_number=int(receiver)).exists() == False:
            messages.error(request, "Receiver's Account Not Found")
            return redirect(reverse('repeattransfer', args=[transfer.pk]))
        if float(amount) < float(100):
            messages.error(request, "Minimum Amount Is ₦100")
            return redirect(reverse('repeattransfer', args=[transfer.pk]))
        if float(amount) > float(100000):
            messages.error(request, "Maximum Amount Is ₦100 000")
            return redirect(reverse('repeattransfer', args=[transfer.pk]))
        if request.user.balance < float(amount):
            messages.error(request, "Insufficient Funds")
            return redirect(reverse('repeattransfer', args=[transfer.pk]))
        receiver_account = CustomerAccount.objects.get(account_number=receiver)
        post = Transfer(sender=obj, txn_id=generate_transfer_id(), receiver=receiver_account, amount=float(amount), purpose=purpose)
        post.save()
        post.user.add(request.user)
        post.user.add(receiver_account.user)
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmtransfer', args=[post.pk]))
    context = {
        'message_c':message_c,
        'obj':obj,
        'transfer':transfer,
        'txn_limit':txn_limit,
        'account_limit':account_limit(request)
        }
    return render(request, 'customer/repeatlocaltransfer.html', context)

@customer_required
def ConfirmTransferView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    transfer = Transfer.objects.get(id=pk)
    if request.method == 'POST':
        messages.info(request, 'Last Step')
        return redirect(reverse('validatetransfer', args=[transfer.pk]))
    context = {
        'message_c':message_c,
        'transfer':transfer
        }
    return render(request, 'customer/confirmtransfer.html', context)

@customer_required
def ValidateTransferView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    transfer = Transfer.objects.get(id=pk)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        if txn_limit < transfer.amount:
            messages.error(request, 'Daily Transaction Limit Has Been Reached')
            return redirect('transfer')
        else:
            pin = request.POST['pin']
            if transfer.sender.user.pin == pin:
                transfer.sender.user.balance -= transfer.amount
                transfer.sender.user.save()
                data = Alert(amount=transfer.amount, balance=transfer.sender.user.balance, user=transfer.sender.user, transfer=transfer, txn='Debit', how='Mobile', which='UTU')
                data.detail = f'{data.transfer.txn_id}/{data.how}/{data.which}/Transfer'
                data.save()
                transfer.receiver.user.balance += transfer.amount
                transfer.receiver.user.save()
                data1 = Alert(amount=transfer.amount, balance=transfer.receiver.user.balance, user=transfer.receiver.user, transfer=transfer, txn='Credit', how='Mobile', which='UTU')
                data1.detail = f'{data.transfer.txn_id}/{data.how}/{data.which}/Transfer'
                data1.save()
                transfer.status = 'Completed'
                transfer.save()
                txn = Txn(account=transfer.sender, transfer=transfer, amount=transfer.amount)
                txn.save()
                if transfer.sender.user.email_notification == True:
                    subject = 'Patridge Bank: Transfer Successful'
                    email_template_name = 'mail/transfer_success.txt'
                    current_site = get_current_site(request)
                    context = {
                    'domain': current_site.domain,
                    "user": request.user,
                    "transfer": transfer,
                    }
                    email = render_to_string(email_template_name, context)
                    try:
                        msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                        msg.send()
                    except BadHeaderError:
                        return redirect('Invalid header found.')
                if transfer.receiver.user.email_notification == True:
                    subject = 'Patridge Bank: Transfer Credit Alert'
                    email_template_name = 'mail/transfer_credit.txt'
                    current_site = get_current_site(request)
                    context = {
                    'domain': current_site.domain,
                    "user": transfer.receiver.user,
                    "transfer": transfer,
                    }
                    email = render_to_string(email_template_name, context)
                    try:
                        msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [transfer.receiver.user.email])
                        msg.send()
                    except BadHeaderError:
                        return redirect('Invalid header found.')
                messages.success(request, f'Transfer Of ₦{transfer.amount} to {transfer.receiver.account_name} Is Successful')
                return redirect('transfer')
            else:
                messages.error(request, 'Invalid Pin')
                return redirect(reverse('validatetransfer', args=[transfer.pk]))
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def BeneficiaryView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    obj = Beneficiary.objects.filter(account=account, confirm=True).all()
    context = {
        'message_c':message_c,
        'obj':obj
        }
    return render(request, 'customer/beneficiary.html', context)

@customer_required
def AddBeneficiaryView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        beneficiary = request.POST['beneficiary']
        if len(beneficiary) > 10 or len(beneficiary) < 10:
            messages.error(request, "Length Of Account Number Should Be 10")
            return redirect('addbeneficiary')
        if CustomerAccount.objects.filter(account_number=int(beneficiary)).exists() == False:
            messages.error(request, "Account Not Found")
            return redirect('addbeneficiary')
        if Beneficiary.objects.filter(account=account, beneficiary=CustomerAccount.objects.get(account_number=beneficiary)).exists():
            messages.error(request, 'Beneficiary Already Added')
            return redirect('addbeneficiary')
        if CustomerAccount.objects.get(account_number=beneficiary) == account:
            messages.error(request, 'You Cannot Add Yourself To Your Beneficiaries')
            return redirect('addbeneficiary')
        else:
            post = Beneficiary(account=account, beneficiary=CustomerAccount.objects.get(account_number=beneficiary))
            post.save()
            messages.info(request, 'Confirm Beneficiary')
            return redirect(reverse('confirmbeneficiary', args=[post.pk]))
    return render(request, 'customer/addbeneficiary.html', {'message_c':message_c})

@customer_required
def ConfirmBeneficiaryView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = Beneficiary.objects.get(id=pk)
    if request.method == 'POST':
        messages.info(request, 'Last Step')
        return redirect(reverse('validatebeneficiary', args=[post.pk]))
    context = {
        'message_c':message_c,
        'beneficiary':post
        }
    return render(request, 'customer/confirmbeneficiary.html', context)

@customer_required
def ValidateBeneficiaryView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    post = Beneficiary.objects.get(id=pk)
    if request.method == 'POST':
        pin = request.POST['pin']
        if post.account.user.pin == pin:
            post.confirm = True
            post.save()
            if post.account.user.email_notification == True:
                subject = 'Patridge Bank: Beneficiary Added Successfully'
                email_template_name = 'mail/beneficiary_success.txt'
                current_site = get_current_site(request)
                context = {
                'domain': current_site.domain,
                "user": request.user,
                "beneficiary": post,
                'date': timezone.now().date()
                }
                email = render_to_string(email_template_name, context)
                try:
                    msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                    msg.send()
                except BadHeaderError:
                    return redirect('Invalid header found.')
            messages.success(request, 'Beneficiary Added Successfully')
            return redirect('beneficiary')
        else:
            messages.error(request, 'Invalid Pin')
            return redirect(reverse('validatebeneficiary', args=[post.pk]))
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def RemoveBeneficiaryView(request):
    if request.method == 'POST':
        beneficiary_id = request.POST.get('beneficiary_id')
        try:
            obj = Beneficiary.objects.get(id=beneficiary_id)
            obj.delete()
            return redirect('beneficiary')
        except Beneficiary.DoesNotExist:
            pass
        except Exception as e:
            pass
    return redirect('beneficiary')

@customer_required
def BeneficiaryTransferView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    txn_limit = account_limit(request) - spent(request)
    beneficiary = Beneficiary.objects.get(id=pk)
    obj = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        receiver = request.POST['receiver']
        amount = request.POST['amount']
        purpose = request.POST['purpose']
        if len(receiver) > 10 or len(receiver) < 10:
            messages.error(request, "Length Of Account Number Should Be 10")
            return redirect(reverse('beneficiarytransfer', args=[beneficiary.pk]))
        if CustomerAccount.objects.filter(account_number=int(receiver)).exists() == False:
            messages.error(request, "Receiver's Account Not Found")
            return redirect(reverse('beneficiarytransfer', args=[beneficiary.pk]))
        if float(amount) < float(100):
            messages.error(request, "Minimum Amount Is ₦100")
            return redirect(reverse('beneficiarytransfer', args=[beneficiary.pk]))
        if float(amount) > float(100000):
            messages.error(request, "Maximum Amount Is ₦100 000")
            return redirect(reverse('beneficiarytransfer', args=[beneficiary.pk]))
        if request.user.balance < float(amount):
            messages.error(request, "Insufficient Funds")
            return redirect(reverse('beneficiarytransfer', args=[beneficiary.pk]))
        receiver_account = CustomerAccount.objects.get(account_number=receiver)
        post = Transfer(sender=obj, txn_id=generate_transfer_id(), receiver=receiver_account, amount=float(amount), purpose=purpose)
        post.save()
        post.user.add(request.user)
        post.user.add(receiver_account.user)
        messages.info(request, 'Confirm Transaction')
        return redirect(reverse('confirmtransfer', args=[post.pk]))
    context = {
        'message_c':message_c,
        'obj':obj,
        'beneficiary':beneficiary,
        'txn_limit':txn_limit,
        'account_limit':account_limit(request)
        }
    return render(request, 'customer/beneficiarytransfer.html', context)

@customer_required
def ServiceView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    if Card.objects.filter(account__id=request.user.id, expires__gt=datetime.datetime.now().date()).exists():
        card = Card.objects.get(account__id=request.user.id, expires__gt=datetime.datetime.now().date())
    else:
        card = False
    return render(request, 'customer/services.html', {'message_c':message_c, 'card':card})

@customer_required
def CardView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    if Card.objects.filter(account__id=request.user.id, expires__gt=datetime.datetime.now().date()).exists():
        card = Card.objects.get(account__id=request.user.id, expires__gt=datetime.datetime.now().date())
    else:
        card = False
    account = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        messages.info(request, 'Your Request Has Been Received')
        return redirect('validatecard')
    return render(request, 'customer/requestcard.html', {'message_c':message_c, 'card':card, 'account':account})

@customer_required
def ValidateCardView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    txn_limit = account_limit(request) - spent(request)
    if request.method == 'POST':
        if txn_limit < float(1000):
            messages.error(request, 'Daily Transaction Limit Has Been Reached')
            return redirect('service')
        else:
            if account.user.balance < float(1000):
                messages.error(request, 'A Minimum Of ₦1000 Is Required For This Request')
                return redirect('service')
            else:
                if Card.objects.filter(account__id=request.user.id, expires__gt=datetime.datetime.now().date()).exists():
                    messages.error(request, 'A Card Is Already Attached To This Account')
                    return redirect('service')
                else:
                    if account.user.pin == request.POST['pin']:
                        if request.POST['pin1'] == request.POST['pin2']:
                            account.user.balance -= float(1000)
                            account.user.save()
                            new = Card(account=account, card_id=generate_card_id(), card_number=generate_card_number(), cvv=generate_cvv(), expires=expires(), card_type='VISA', card='Debit', pin=request.POST['pin1'])
                            try:
                                new.save()
                            except(ValueError):
                                messages.error(request, 'Pin Should Be Digits')
                                return redirect('validatecard')
                            alert = Alert(amount=float(1000), balance=account.user.balance, user=request.user, card=new, txn='Debit', how='Card', which='Patridge Bank')
                            alert.detail = f'{alert.card.card_id}/{alert.which}/{alert.how} Payment'
                            alert.save()
                            txn = Txn(account=account, amount=float(1000))
                            txn.save()
                            message = Message(user=account.user, message=f'Patridge Bank: Congrats, You Have Gotten A Patridge Bank Virtual Card')
                            message.save()
                            if account.user.email_notification == True:
                                subject = 'Patridge Bank: Virtual Card Request Successful'
                                email_template_name = 'mail/card_success.txt'
                                current_site = get_current_site(request)
                                context = {
                                'domain': current_site.domain,
                                "user": request.user,
                                "card": new,
                                }
                                email = render_to_string(email_template_name, context)
                                try:
                                    msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                                    msg.send()
                                except BadHeaderError:
                                    return redirect('Invalid header found.')
                            messages.success(request, 'Congrats, You Have Gotten A Patridge Bank Virtual Card')
                            return redirect('service')
                        else:
                            messages.error(request, 'Card Pin Not Thesame')
                            return redirect('validatecard')
                    else:
                        messages.error(request, 'Invalid Transaction Pin')
                        return redirect('validatecard')
    return render(request, 'customer/validatecard.html', {'message_c':message_c})

@customer_required
def ModifyLimitView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        if 'agree' in request.POST:
            account.bvn = request.POST['bvn']
            account.save()
            messages.success(request, 'Almost There')
            return redirect('validatemodify')
        else:
            messages.error(request, 'You Have To Agree To Our Terms And Policies To Continue')
            return redirect('modify')
    return render(request, 'customer/modifylimit.html', {'message_c':message_c, 'account':account})

@customer_required
def ValidateModifyView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        if account.user.balance < float(1000):
            messages.error(request, 'A Minimum Of ₦1000 Is Required For This Request')
            return redirect('service')
        else:
            if account.txn_limit == float(1000000):
                messages.error(request, 'Your Daily Transaction Limit Has Already Been Upgraded And Cannot Be Further Modified')
                return redirect('service')
            else:
                if account.user.pin == request.POST['pin']:
                    account.user.balance -= float(1000)
                    account.user.save()
                    account.txn_limit = float(1000000)
                    account.save()
                    alert = Alert(amount=float(1000), balance=account.user.balance, user=request.user, modify_limit=True, txn='Debit', how='Mobile', which='Patridge Bank')
                    alert.detail = f'Modify_Transfer_Limit/{alert.how}/{alert.which}'
                    alert.save()
                    message = Message(user=account.user, message=f'Patridge Bank: Congrats, Your Daily Transaction Limit Has Been Upgraded')
                    message.save()
                    if account.user.email_notification == True:
                        subject = 'Patridge Bank: Daily Transaction Limit Upgraded'
                        email_template_name = 'mail/limit_success.txt'
                        current_site = get_current_site(request)
                        context = {
                        'domain': current_site.domain,
                        "user": request.user,
                        "account": account,
                        'date': timezone.now().date()
                        }
                        email = render_to_string(email_template_name, context)
                        try:
                            msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                            msg.send()
                        except BadHeaderError:
                            return redirect('Invalid header found.')
                    messages.success(request, 'Congrats, Your Daily Transaction Limit Has Been Upgraded')
                    return redirect('service')
                else:
                    messages.error(request, 'Invalid Pin')
                    return redirect('validatemodify')
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def StatementView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    alert = Alert.objects.filter(user_id=request.user.id)
    if request.method == 'POST':
        start = request.POST['start']
        stop = request.POST['stop']
        if not 'doc' in request.POST and not 'email' in request.POST:
            messages.error(request, 'Choose A Method')
            return redirect('statement')
        if 'doc' in request.POST and 'email' in request.POST:
            messages.error(request, 'Choose Only One Method')
            return redirect('statement')
        if request.POST['start'] > request.POST['stop']:
            messages.error(request, 'Start Date Should Not Be Greater Than Stop Date')
            return redirect('statement')
        if request.POST['start'] > str(timezone.now().date()) or request.POST['stop'] > str(timezone.now().date()):
            messages.error(request, 'Date Should Not Be More Than Today')
            return redirect('statement')
        if account.user.balance < float(100):
            messages.error(request, 'Insufficient Funds')
            return redirect('service')
        if 'doc' in request.POST and not 'email' in request.POST:
            if account.user.pin == request.POST['pin']:
                if alert.filter(user=request.user, date__date__range=(start, stop)).exists():
                    account.user.balance -= float(100)
                    account.user.save()
                    statement = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop))
                    doc = Document()
                    doc.add_heading('Patridge Bank', level=1)
                    doc.add_heading('Bank Statement', level=2)
                    doc.add_paragraph(f'Account Name: {account.account_name}')
                    doc.add_paragraph(f'Account Number: {account.account_number}')
                    doc.add_paragraph(f'Date: {timezone.now().date()}')
                    doc.add_paragraph(f'Statement From: {start}')
                    doc.add_paragraph(f'Statement To: {stop}')
                    if Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Debit').exists():
                        debits = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Debit')
                        trans = [float(debit.amount) for debit in debits]
                        total_debit = sum(trans)
                    else:
                        total_debit = 0
                    if Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Credit').exists():
                        credits = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Credit')
                        trans = [float(credit.amount) for credit in credits]
                        total_credit = sum(trans)
                    else:
                        total_credit = 0
                    doc.add_paragraph(f'Total Credits: ₦{total_credit}')
                    doc.add_paragraph(f'Total Debits: ₦{total_debit}')
                    doc.add_paragraph(f'Latest Balance: ₦{account.user.balance}')
                    doc.add_heading('Statement Details', level=1)
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    table.cell(0,0).text = 'Date'
                    table.cell(0,1).text = 'Transaction'
                    table.cell(0,2).text = 'Details'
                    table.cell(0,3).text = 'Amount(₦)'
                    table.cell(0,4).text = 'Balance(₦)'
                    for obj in statement:
                        row = table.add_row().cells
                        row[0].text = str(obj.date)
                        row[1].text = str(obj.txn)
                        row[2].text = str(obj.detail)
                        row[3].text = str(obj.amount)
                        row[4].text = str(obj.balance)
                    doc_io = BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    response = FileResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={account.account_name}/statement/{timezone.now()}.docx'
                    new = Alert(amount=float(100), balance=account.user.balance, user=request.user, statement=True, txn='Debit', how='Mobile', which='Patridge Bank')
                    new.detail = f'Bank Statement/{new.how}/{new.which}'
                    new.save()
                    return response
                else:
                    messages.error(request, 'No Transaction Occurred Within The Date Range Specified')
                    return redirect('statement')
            else:
                messages.error(request, 'Invalid Pin')
                return redirect('statement')
        if not 'doc' in request.POST and 'email' in request.POST:
            if account.user.pin == request.POST['pin']:
                if alert.filter(user=request.user, date__date__range=(start, stop)).exists():
                    account.user.balance -= float(100)
                    account.user.save()
                    subject = 'Patridge Bank: Bank Statement Request'
                    email_template_name = 'mail/statement_success.txt'
                    current_site = get_current_site(request)
                    context = {
                    'domain': current_site.domain,
                    "user": request.user,
                    "account": account,
                    "start": start,
                    "stop": stop,
                    'date': timezone.now().date()
                    }
                    email = render_to_string(email_template_name, context)
                    try:
                        msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                        msg.send()
                    except BadHeaderError:
                        return redirect('Invalid header found.')
                    new = Alert(amount=float(100), balance=account.user.balance, user=request.user, statement=True, txn='Debit', how='Mobile', which='Patridge Bank')
                    new.detail = f'Bank Statement/{new.how}/{new.which}'
                    new.save()
                    messages.success(request, "Statement Request Successful. Check Your Mail's Inbox Or Spam Folder")
                    return redirect('service')
                else:
                    messages.error(request, 'No Transaction Occurred Within The Date Range Specified')
                    return redirect('statement')
            else:
                messages.error(request, 'Invalid Pin')
                return redirect('statement')
    return render(request, 'customer/requeststatement.html', {'message_c':message_c, 'account':account})

@customer_required
def DownloadStatementView(request, start, stop):
    account = CustomerAccount.objects.get(id=request.user.id)
    statement = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop))
    doc = Document()
    doc.add_heading('Patridge Bank', level=1)
    doc.add_heading('Bank Statement', level=2)
    doc.add_paragraph(f'Account Name: {account.account_name}')
    doc.add_paragraph(f'Account Number: {account.account_number}')
    doc.add_paragraph(f'Date: {timezone.now().date()}')
    doc.add_paragraph(f'Statement From: {start}')
    doc.add_paragraph(f'Statement To: {stop}')
    if Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Debit').exists():
        debits = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Debit')
        trans = [float(debit.amount) for debit in debits]
        total_debit = sum(trans)
    else:
        total_debit = 0
    if Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Credit').exists():
        credits = Alert.objects.filter(user_id=request.user.id, date__date__range=(start, stop), txn='Credit')
        trans = [float(credit.amount) for credit in credits]
        total_credit = sum(trans)
    else:
        total_credit = 0
    doc.add_paragraph(f'Total Credits: ₦{total_credit}')
    doc.add_paragraph(f'Total Debits: ₦{total_debit}')
    doc.add_paragraph(f'Latest Balance: ₦{account.user.balance}')
    doc.add_heading('Statement Details', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.cell(0,0).text = 'Date'
    table.cell(0,1).text = 'Transaction'
    table.cell(0,2).text = 'Details'
    table.cell(0,3).text = 'Amount(₦)'
    table.cell(0,4).text = 'Balance(₦)'
    for obj in statement:
        row = table.add_row().cells
        row[0].text = str(obj.date)
        row[1].text = str(obj.txn)
        row[2].text = str(obj.detail)
        row[3].text = str(obj.amount)
        row[4].text = str(obj.balance)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    response = FileResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={account.account_name}/statement/{timezone.now()}.docx'
    return response

@customer_required
def LocateView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    q = request.POST.get('q') if request.POST.get('q') != None else ''
    obj = Branch.objects.filter( Q(name__icontains=q) | Q(address__icontains=q) | Q(state__icontains=q)).order_by('name')
    obj_c = Branch.objects.filter( Q(name__icontains=q) | Q(address__icontains=q) | Q(state__icontains=q)).order_by('name').count()
    context={'message_c':message_c, 'branch':obj, 'branch_c':obj_c}
    return render(request, 'customer/locate.html', context)

@customer_required
def ContactView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    return render(request, 'customer/contact.html', {'message_c':message_c})

@customer_required
def ProfileView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(id=request.user.id)
    if request.method == 'POST':
        form = PictureForm(request.POST, request.FILES)
        if form.is_valid():
            account.picture = form.cleaned_data['picture']
            account.save()
            messages.success(request, 'Picture Added Successfully')
            return redirect('profile')
    else:
        form = PictureForm()
    return render(request, 'customer/profile.html', {'message_c':message_c, 'account':account, 'form':form})

@customer_required
def SettingView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    if request.method == 'POST':
        if request.user.email_notification == True:
            request.user.email_notification = False
            request.user.save()
            messages.success(request, 'Email Notification Is Off')
            return redirect('setting')
        if request.user.email_notification == False:
            request.user.email_notification = True
            request.user.save()
            messages.success(request, 'Email Notification Is On')
            return redirect('setting')
    return render(request, 'customer/setting.html', {'message_c':message_c, 'user':request.user})

@customer_required
def ChangeUserIDView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    if request.method == 'POST':
        if request.POST['old_user_id'] != account.user.user_ID:
            messages.error(request, 'Wrong User ID')
            return redirect('changeuserid')
        if len(request.POST['new_user_id']) > 11:
            messages.error(request, 'User ID Should Not Be More Than 11 Digits')
            return redirect('changeuserid')
        if request.POST['pin'] != account.user.pin:
            messages.error(request, 'Invalid Pin')
            return redirect('changeuserid')
        else:
            account.user.user_ID = request.POST['new_user_id']
            account.user.save()
            message = Message(user=account.user, message=f'Patridge Bank: User ID Changed Successfully')
            message.save()
            if account.user.email_notification == True:
                subject = 'Patridge Bank: User ID Change Success'
                email_template_name = 'mail/user_id_success.txt'
                current_site = get_current_site(request)
                context = {
                'domain': current_site.domain,
                "user": request.user,
                "account": account,
                'date': timezone.now().date()
                }
                email = render_to_string(email_template_name, context)
                try:
                    msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                    msg.send()
                except BadHeaderError:
                    return redirect('Invalid header found.')
            messages.success(request, 'User ID Changed Successfully')
            return redirect('setting')
    return render(request, 'customer/changeuserid.html', {'message_c':message_c})

@customer_required
def ChangePinView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    if request.method == 'POST':
        if account.user.pin == request.POST['pin']:
            messages.success(request, 'Enter New Pin')
            return redirect('confirmpin')
        else:
            messages.error(request, 'Invalid Pin')
            return redirect('changepin')
    else:
        messages.info(request, 'Enter Old Pin')
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def ConfirmPinView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    if request.method == 'POST':
        pin = request.POST['pin']
        post = User.objects.filter(pin=pin)
        if post.exists():
            messages.error(request, 'Pin Is Too Common')
            return redirect('confirmpin')
        if pin == '0000' or pin == '1234' or pin == '1111' or pin == '2222' or pin == '3333' or pin == '4444' or pin == '5555' or pin == '6666' or pin == '7777' or pin == '8888' or pin == '9999':
            messages.error(request, 'Pin Is Too Common')
            return redirect('confirmpin')
        else:
            account.user.pin = pin
            account.user.save()
            message = Message(user=account.user, message=f'Patridge Bank: Transaction Pin Changed Successfully')
            message.save()
            if account.user.email_notification == True:
                subject = 'Patridge Bank: Transaction Pin Change Success'
                email_template_name = 'mail/pin_change_success.txt'
                current_site = get_current_site(request)
                context = {
                'domain': current_site.domain,
                "user": request.user,
                "account": account,
                'date': timezone.now().date()
                }
                email = render_to_string(email_template_name, context)
                try:
                    msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                    msg.send()
                except BadHeaderError:
                    return redirect('Invalid header found.')
            messages.success(request, 'Transaction Pin Changed Successfully')
            return redirect('setting')
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def SupportView(request):
    try:
        Support.objects.read_support(customer_id=request.user.id)
    except (Support.DoesNotExist, UnboundLocalError):
        pass
    message_c = Message.objects.filter(user=request.user, read=False).count()
    support = Support.objects.filter(customer=request.user).order_by('date')
    paginator = Paginator(support, 10)
    last_page = paginator.num_pages
    page = request.GET.get('page', last_page)
    try:
        support = paginator.page(page)
    except PageNotAnInteger:
        support = paginator.page(1)
    except EmptyPage:
        support = paginator.page(paginator.num_pages)
    if request.method == 'POST':
        message = request.POST['message']
        new = Support(customer=request.user, message=message)
        new.save()
        auto = Support(answer=True, customer=request.user, staff=request.user, message='Patridge Bank Has Received Your Support Message; Expect A Response Within 24 Hours')
        auto.save()
        return redirect('support')
    return render(request, 'customer/support.html', {'message_c':message_c, 'support':support})

@customer_required
def FAQView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    faq = FAQ.objects.all()
    return render(request, 'customer/faq.html', {'message_c':message_c, 'faq':faq})

@customer_required
def LoanView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    loan = Loan.objects.filter(account=account).order_by('-date')
    if Loan.objects.filter(account=account, approved=True, paid=False, till__lt=timezone.now().date()).exists():
        in_debt_object = Loan.objects.get(account=account, approved=True, paid=False, till__lt=timezone.now().date())
        in_debt = True
        interest = (0.5/100) * in_debt_object.amount
        if timezone.now().date() - in_debt_object.date.date():
            day = timezone.now().date() - in_debt_object.date.date()
            days = day.days
        else:
            days = 0
        total_interest = float(interest) * float(days)
        in_debt_object.new_amount = in_debt_object.amount + total_interest
        in_debt_object.save()
    else:
        in_debt_object = None
        in_debt = False
    if Loan.objects.filter(account=account, approved=False, disapproved=False).exists():
        requested = True
    else:
        requested = False
    if Loan.objects.filter(account=account, approved=True, paid=False, till__gte=timezone.now().date()).exists():
        valid_loan_object = Loan.objects.get(account=account, approved=True, paid=False, till__gte=timezone.now().date())
        valid_loan = True
        interest = (0.5/100) * valid_loan_object.amount
        if timezone.now().date() - valid_loan_object.date.date():
            day = timezone.now().date() - valid_loan_object.date.date()
            days = day.days
        else:
            days = 0
        total_interest = float(interest) * float(days)
        valid_loan_object.new_amount = valid_loan_object.amount + total_interest
        valid_loan_object.save()
    else:
        valid_loan_object = None
        valid_loan = False
    if not Loan.objects.filter(account=account, approved=True, paid=False).exists():
        no_debt = True
        if date_eligible(request).date() < timezone.now().date() and transaction_eligible(request) > float(10000):
            date_eligibility = True
            transaction_eligibility = True
        else:
            date_eligibility = False
            transaction_eligibility = False
    else:
        no_debt = False
        date_eligibility = False
        transaction_eligibility = False
    if request.method == 'POST':
        if not Loan.objects.filter(account=account, approved=True, paid=False).exists():
            if not Loan.objects.filter(account=account, approved=False, disapproved=False):
                if '1week' in request.POST and not '1month' in request.POST and not '3months' in request.POST and not '6months' in request.POST:
                    if account.user.pin == request.POST['pin']:
                        new = Loan(account=account, reason=request.POST['reason'], amount=float(request.POST['amount']), new_amount=float(request.POST['amount']), source_of_income=request.POST['income'], period='1 Week')
                        new.save()
                        message = Message(user=account.user, message=f'Patridge Bank: Loan Request Successful')
                        message.save()
                        if account.user.email_notification == True:
                            subject = 'Patridge Bank: Loan Request Successful'
                            email_template_name = 'mail/loan_request_success.txt'
                            current_site = get_current_site(request)
                            context = {
                            'domain': current_site.domain,
                            "user": request.user,
                            'new':new,
                            "account": account,
                            'date': timezone.now().date()
                            }
                            email = render_to_string(email_template_name, context)
                            try:
                                msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                                msg.send()
                            except BadHeaderError:
                                return redirect('Invalid header found.')
                        messages.success(request, 'Loan Request Successful')
                        return redirect('loan')
                    else:
                        messages.error(request, 'Invalid Pin')
                        return redirect('loan')
                elif not '1week' in request.POST and '1month' in request.POST and not '3months' in request.POST and not '6months' in request.POST:
                    if account.user.pin == request.POST['pin']:
                        new = Loan(account=account, reason=request.POST['reason'], amount=float(request.POST['amount']), new_amount=float(request.POST['amount']), source_of_income=request.POST['income'], period='1 Month')
                        new.save()
                        message = Message(user=account.user, message=f'Patridge Bank: Loan Request Successful')
                        message.save()
                        if account.user.email_notification == True:
                            subject = 'Patridge Bank: Loan Request Successful'
                            email_template_name = 'mail/loan_request_success.txt'
                            current_site = get_current_site(request)
                            context = {
                            'domain': current_site.domain,
                            "user": request.user,
                            'new':new,
                            "account": account,
                            'date': timezone.now().date()
                            }
                            email = render_to_string(email_template_name, context)
                            try:
                                msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                                msg.send()
                            except BadHeaderError:
                                return redirect('Invalid header found.')
                        messages.success(request, 'Loan Request Successful')
                        return redirect('loan')
                    else:
                        messages.error(request, 'Invalid Pin')
                        return redirect('loan')
                elif not '1week' in request.POST and not '1month' in request.POST and '3months' in request.POST and not '6months' in request.POST:
                    if account.user.pin == request.POST['pin']:
                        new = Loan(account=account, reason=request.POST['reason'], amount=float(request.POST['amount']), new_amount=float(request.POST['amount']), source_of_income=request.POST['income'], period='3 Months')
                        new.save()           
                        message = Message(user=account.user, message=f'Patridge Bank: Loan Request Successful')
                        message.save()
                        if account.user.email_notification == True:
                            subject = 'Patridge Bank: Loan Request Successful'
                            email_template_name = 'mail/loan_request_success.txt'
                            current_site = get_current_site(request)
                            context = {
                            'domain': current_site.domain,
                            "user": request.user,
                            'new':new,
                            "account": account,
                            'date': timezone.now().date()
                            }
                            email = render_to_string(email_template_name, context)
                            try:
                                msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                                msg.send()
                            except BadHeaderError:
                                return redirect('Invalid header found.')
                        messages.success(request, 'Loan Request Successful')
                        return redirect('loan')
                    else:
                        messages.error(request, 'Invalid Pin')
                        return redirect('loan')
                elif not '1week' in request.POST and not '1month' in request.POST and not '3months' in request.POST and '6months' in request.POST:
                    if account.user.pin == request.POST['pin']:
                        new = Loan(account=account, reason=request.POST['reason'], amount=float(request.POST['amount']), new_amount=float(request.POST['amount']), source_of_income=request.POST['income'], period='6 Months')
                        new.save()
                        message = Message(user=account.user, message=f'Patridge Bank: Loan Request Successful')
                        message.save()
                        if account.user.email_notification == True:
                            subject = 'Patridge Bank: Loan Request Successful'
                            email_template_name = 'mail/loan_request_success.txt'
                            current_site = get_current_site(request)
                            context = {
                            'domain': current_site.domain,
                            "user": request.user,
                            'new':new,
                            "account": account,
                            'date': timezone.now().date()
                            }
                            email = render_to_string(email_template_name, context)
                            try:
                                msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                                msg.send()
                            except BadHeaderError:
                                return redirect('Invalid header found.')
                        messages.success(request, 'Loan Request Successful')
                        return redirect('loan')
                    else:
                        messages.error(request, 'Invalid Pin')
                        return redirect('loan')
                else:
                    messages.error(request, 'Invalid Choice In Date')
                    return redirect('loan')  
            else:
                messages.error(request, 'You Have Already Made A Request That Has Not Been Responded To')
                return redirect('loan')       
        else:
            messages.error(request, 'You Already Have An Outstanding Loan; Pay Up Then You Can Make Another Request')
            return redirect('loan')
    context = {
        'account':account,
        'message_c':message_c,
        'valid_loan':valid_loan,
        'valid_loan_object':valid_loan_object,
        'no_debt':no_debt,
        'in_debt':in_debt,
        'in_debt_object':in_debt_object,
        'requested':requested,
        'loan':loan,
        'date_eligibility':date_eligibility,
        'transaction_eligibility':transaction_eligibility
        }
    return render(request, 'customer/loan.html', context)

@customer_required
def LoanDetailView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    loan = Loan.objects.get(id=pk)
    context = {
        'message_c':message_c,
        'loan':loan
    }
    return render(request, 'customer/loandetail.html', context)

@customer_required
def CalculatorView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    if request.method == 'POST':
        if 'principal' in request.POST and 'time' in request.POST:
            principal = float(request.POST['principal'])
            time = float(request.POST['time'])
            interest = ((0.5/100) * principal) * time
            total = float(interest) + float(request.POST['principal'])
            context = {
                'message_c':message_c,
                'result':'result',
                'principal':principal,
                'time':time,
                'interest':interest,
                'total':total
                }
            return render(request, 'customer/interestcalculator.html', context)
        else:
            messages.error(request, 'Invalid Input')
            return redirect('calculate')
    return render(request, 'customer/interestcalculator.html', {'message_c':message_c})

@customer_required
def PayLoanView(request, pk):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    loan = Loan.objects.get(id=pk)
    if request.method == 'POST':
        if account.user.pin == request.POST['pin']:
            if account.user.balance > loan.new_amount or account.user.balance == loan.new_amount:
                account.user.balance -= loan.new_amount
                account.user.save()
                loan.paid = True
                loan.save()
                message = Message(user=account.user, message=f'Patridge Bank: Loan Payment Successful')
                message.save()
                alert = Alert(amount=loan.new_amount, balance=account.user.balance, user=request.user, loan=loan, txn='Debit', how='Mobile', which='Patridge Bank')
                alert.detail = f'Loan_Payment/{alert.how}/{alert.which}'
                alert.save()
                if account.user.email_notification == True:
                    subject = 'Patridge Bank: Loan Payment Successful'
                    email_template_name = 'mail/loan_repay_success.txt'
                    current_site = get_current_site(request)
                    context = {
                    'domain': current_site.domain,
                    "user": request.user,
                    'loan':loan,
                    "account": account,
                    'date': timezone.now().date()
                    }
                    email = render_to_string(email_template_name, context)
                    try:
                        msg = EmailMultiAlternatives(subject, email, settings.EMAIL_HOST_USER , [request.user.email])
                        msg.send()
                    except BadHeaderError:
                        return redirect('Invalid header found.')
                messages.success(request, 'Loan Paid Successfully. Thanks For Banking With Us')
                return redirect('loan')
            else:
                messages.error(request, 'Insufficient Funds')
                return redirect('loan')
        else:
            messages.success(request, 'Invalid Pin')
            return redirect(reverse('payloan', args=[loan.pk]))
    return render(request, 'customer/validate.html', {'message_c':message_c})

@customer_required
def MessageView(request):
    message = Message.objects.filter(user=request.user).order_by('-date')[:10]
    message_c = Message.objects.filter(user=request.user, read=False).count()
    context = {
        'message':message,
        'message_c':message_c,
        }
    return render(request, 'customer/message.html', context)

@customer_required
def DetailMessageView(request, pk):
    message = Message.objects.read_message(message_id=pk)
    return redirect('message')

@customer_required
def FundWalletView(request):
    message_c = Message.objects.filter(user=request.user, read=False).count()
    account = CustomerAccount.objects.get(user=request.user)
    form = FundForm(request.POST or None)
    if request.method == 'POST':
        if form.is_valid():
            data = form.save(commit=False)
            data.account = account
            data.save()
            pk = settings.PAYSTACK_PUBLIC_KEY
            context = {
                'message_c':message_c,
                'data': data,
                'account': account,
                'field_values': request.POST,
                'paystack_pub_key': pk,
                'price_value': data.price_value(),
            }
            return render(request, 'customer/payment.html', context)
    return render(request, 'customer/fund.html', {'account': account, "form":form, 'message_c':message_c})            

@customer_required
def VerifyPaymentView(request, ref):
    payment = Wallet.objects.get(ref=ref)
    verified = payment.verify_payment()
    if verified:
        payment.account.user.balance += payment.amount
        payment.account.user.save()
        data = Alert(amount=payment.amount, balance=payment.account.user.balance, user=payment.account.user, txn='Credit', how='Mobile', which='Patridge Bank')
        data.detail = f'FundWallet/{data.how}/{data.which}'
        data.save()
        messages.success(request, f'Your Account Has Been Credited With ₦{payment.amount}')
        return redirect('customer_home')
    else:
        messages.error(request, 'Account Funding Failed')
        return redirect('customer_home')
